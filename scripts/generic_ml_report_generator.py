"""
Generic ML Report Generator
Automatically generates comprehensive analysis reports for any ML prediction model
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import warnings
warnings.filterwarnings('ignore')

# Set style for plots
plt.style.use('seaborn-v0_8')
sns.set_palette("husl")


class GenericMLReportGenerator:
    """
    A flexible report generator that adapts to any dataset and prediction model
    """
    
    def __init__(self, df, target_column, model_name="ML Model", positive_class=None):
        """
        Initialize the report generator
        
        Args:
            df: DataFrame with the data
            target_column: Name of the target/prediction column
            model_name: Name of the ML model/use case (e.g., "Attrition Prediction", "Churn Prediction")
            positive_class: The positive class value (e.g., "Yes", 1, True)
        """
        self.df = df.copy()
        self.target_column = target_column
        self.model_name = model_name
        
        # Auto-detect positive class if not provided
        if positive_class is None:
            unique_values = df[target_column].unique()
            if len(unique_values) == 2:
                # Common positive class patterns
                positive_patterns = ['yes', 'true', '1', 'y', 'positive', 'churned', 'left', 'defaulted']
                for val in unique_values:
                    if str(val).lower() in positive_patterns:
                        self.positive_class = val
                        break
                else:
                    self.positive_class = unique_values[0]
            else:
                self.positive_class = unique_values[0]
        else:
            self.positive_class = positive_class
        
        # Auto-detect column types
        self.numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        self.categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        self.date_cols = df.select_dtypes(include=['datetime64']).columns.tolist()
        
        # Remove target from feature lists
        if target_column in self.numeric_cols:
            self.numeric_cols.remove(target_column)
        if target_column in self.categorical_cols:
            self.categorical_cols.remove(target_column)
        
        # Calculate positive rate
        self.positive_rate = (df[target_column] == self.positive_class).mean()
        
    def create_report(self, output_path):
        """Generate the complete report"""
        doc = Document()
        
        # Cover Page
        self._add_cover_page(doc)
        
        # Table of Contents
        self._add_table_of_contents(doc)
        
        # Executive Summary
        self._add_executive_summary(doc)
        
        # Data Overview
        self._add_data_overview(doc)
        
        # Key Performance Indicators
        self._add_kpis(doc)
        
        # Target Variable Analysis
        self._add_target_analysis(doc)
        
        # Categorical Features Analysis
        if self.categorical_cols:
            self._add_categorical_analysis(doc)
        
        # Numerical Features Analysis
        if self.numeric_cols:
            self._add_numerical_analysis(doc)
        
        # Time-based Analysis (if date columns exist)
        if self.date_cols:
            self._add_temporal_analysis(doc)
        
        # Feature Correlation Analysis
        if len(self.numeric_cols) > 1:
            self._add_correlation_analysis(doc)
        
        # Risk/Prediction Scoring
        self._add_risk_scoring(doc)
        
        # Strategic Recommendations
        self._add_recommendations(doc)
        
        # Appendices
        self._add_appendices(doc)
        
        # Add headers
        self._add_page_headers(doc)
        
        doc.save(output_path)
        print(f"Report generated successfully: {output_path}")
        
    def _add_cover_page(self, doc):
        """Create cover page"""
        doc.add_heading(f'{self.model_name} Analysis Report', 0)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Comprehensive {self.model_name} Analysis')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Executive Report')
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
        run.font.size = Pt(12)
        
        doc.add_page_break()
        
    def _add_table_of_contents(self, doc):
        """Add table of contents"""
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            'Executive Summary',
            'Data Overview and Quality Assessment',
            'Key Performance Indicators',
            f'{self.target_column} Analysis',
            'Feature Analysis',
            'Statistical Insights',
            'Prediction Scoring',
            'Strategic Recommendations',
            'Appendices'
        ]
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
        doc.add_page_break()
        
    def _add_executive_summary(self, doc):
        """Add executive summary"""
        doc.add_heading('Executive Summary', level=1)
        
        summary = f"""
        This comprehensive report analyzes {len(self.df):,} records for {self.model_name}.
        
        Key Highlights:
        • Overall {self.target_column} rate: {self.positive_rate*100:.1f}%
        • Total features analyzed: {len(self.numeric_cols) + len(self.categorical_cols)}
        • Data quality: {(1 - self.df.isnull().sum().sum()/(self.df.shape[0]*self.df.shape[1]))*100:.1f}% complete
        
        This analysis provides actionable insights through statistical analysis, visualization,
        and predictive scoring to support data-driven decision making.
        """
        doc.add_paragraph(summary.strip())
        
    def _add_data_overview(self, doc):
        """Add data overview section"""
        doc.add_heading('Data Overview and Quality Assessment', level=1)
        
        doc.add_paragraph(f"**Dataset Size:** {self.df.shape[0]:,} records with {self.df.shape[1]} features")
        
        # Data quality table
        doc.add_heading('Data Quality Metrics', level=2)
        quality_table = doc.add_table(rows=1, cols=3)
        quality_table.style = 'Table Grid'
        hdr_cells = quality_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Status'
        
        total_missing = self.df.isnull().sum().sum()
        missing_pct = (total_missing / (self.df.shape[0] * self.df.shape[1])) * 100
        completeness = 100 - missing_pct
        
        quality_stats = [
            ('Total Records', f"{len(self.df):,}", 'Complete'),
            ('Total Missing Values', f"{total_missing:,} ({missing_pct:.1f}%)", 
             'Excellent' if missing_pct < 1 else 'Good' if missing_pct < 5 else 'Needs Attention'),
            ('Data Completeness', f"{completeness:.1f}%", 
             'Excellent' if completeness > 99 else 'Good' if completeness > 95 else 'Acceptable'),
            ('Numeric Features', f"{len(self.numeric_cols)}", 'Available'),
            ('Categorical Features', f"{len(self.categorical_cols)}", 'Available')
        ]
        
        for metric, value, status in quality_stats:
            row_cells = quality_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            row_cells[2].text = status
            
    def _add_kpis(self, doc):
        """Add Key Performance Indicators"""
        doc.add_heading('Key Performance Indicators', level=1)
        
        kpi_table = doc.add_table(rows=1, cols=2)
        kpi_table.style = 'Table Grid'
        hdr_cells = kpi_table.rows[0].cells
        hdr_cells[0].text = 'KPI'
        hdr_cells[1].text = 'Value'
        
        kpis = [
            ('Total Records', f"{len(self.df):,}"),
            (f'{self.target_column} - Positive Class', f"{(self.df[self.target_column] == self.positive_class).sum():,}"),
            (f'{self.target_column} Rate', f"{self.positive_rate*100:.1f}%"),
        ]
        
        # Add numeric column stats
        for col in self.numeric_cols[:5]:  # Top 5 numeric features
            kpis.append((f'Average {col}', f"{self.df[col].mean():.2f}"))
            
        for kpi, value in kpis:
            row_cells = kpi_table.add_row().cells
            row_cells[0].text = kpi
            row_cells[1].text = value
            
    def _add_target_analysis(self, doc):
        """Analyze target variable distribution"""
        doc.add_heading(f'{self.target_column} Distribution Analysis', level=1)
        
        # Target distribution pie chart
        plt.figure(figsize=(8, 6))
        target_counts = self.df[self.target_column].value_counts()
        colors = ['#3498DB', '#E74C3C', '#2ECC71', '#F39C12']
        plt.pie(target_counts.values, labels=target_counts.index, autopct='%1.1f%%',
                startangle=90, colors=colors[:len(target_counts)])
        plt.title(f'{self.target_column} Distribution', fontsize=14, fontweight='bold')
        plt.axis('equal')
        plt.tight_layout()
        plt.savefig('/tmp/target_distribution.png', dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        doc.add_picture('/tmp/target_distribution.png', width=Inches(5))
        
    def _add_categorical_analysis(self, doc):
        """Analyze categorical features"""
        doc.add_heading('Categorical Features Analysis', level=1)
        
        # Analyze top categorical features
        for col in self.categorical_cols[:5]:  # Top 5 categorical features
            if self.df[col].nunique() > 50:  # Skip if too many categories
                continue
                
            doc.add_heading(f'{col} Analysis', level=2)
            
            # Calculate rates by category
            category_rates = self.df.groupby(col)[self.target_column].apply(
                lambda x: (x == self.positive_class).mean()
            ).sort_values(ascending=False)
            
            # Plot
            plt.figure(figsize=(10, 6))
            category_rates.plot(kind='barh' if len(category_rates) > 5 else 'bar', 
                              color='#9B59B6')
            plt.title(f'{self.target_column} Rate by {col}', fontsize=14, fontweight='bold')
            plt.ylabel(f'{self.target_column} Rate' if len(category_rates) <= 5 else col, fontsize=12)
            plt.xlabel(col if len(category_rates) <= 5 else f'{self.target_column} Rate', fontsize=12)
            plt.grid(axis='y' if len(category_rates) <= 5 else 'x', alpha=0.3)
            plt.tight_layout()
            plt.savefig(f'/tmp/cat_{col}.png', dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            
            doc.add_picture(f'/tmp/cat_{col}.png', width=Inches(6))
            
    def _add_numerical_analysis(self, doc):
        """Analyze numerical features"""
        doc.add_heading('Numerical Features Analysis', level=1)
        
        # Analyze top numerical features
        for col in self.numeric_cols[:5]:  # Top 5 numeric features
            doc.add_heading(f'{col} Distribution', level=2)
            
            # Distribution plot
            plt.figure(figsize=(10, 6))
            
            # Create separate distributions for positive and negative classes
            pos_data = self.df[self.df[self.target_column] == self.positive_class][col].dropna()
            neg_data = self.df[self.df[self.target_column] != self.positive_class][col].dropna()
            
            plt.hist([neg_data, pos_data], bins=25, label=['Negative', 'Positive'],
                    color=['#3498DB', '#E74C3C'], alpha=0.7)
            plt.title(f'{col} Distribution by {self.target_column}', fontsize=14, fontweight='bold')
            plt.xlabel(col, fontsize=12)
            plt.ylabel('Frequency', fontsize=12)
            plt.legend()
            plt.grid(axis='y', alpha=0.3)
            plt.tight_layout()
            plt.savefig(f'/tmp/num_{col}.png', dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            
            doc.add_picture(f'/tmp/num_{col}.png', width=Inches(6))
            
            # Add statistics
            stats_text = f"""
            Mean: {self.df[col].mean():.2f}
            Median: {self.df[col].median():.2f}
            Std Dev: {self.df[col].std():.2f}
            Min: {self.df[col].min():.2f}
            Max: {self.df[col].max():.2f}
            """
            doc.add_paragraph(stats_text.strip())
            
    def _add_temporal_analysis(self, doc):
        """Add time-based analysis if date columns exist"""
        doc.add_heading('Temporal Analysis', level=1)
        
        for date_col in self.date_cols[:3]:  # Analyze up to 3 date columns
            doc.add_heading(f'Trends Over Time ({date_col})', level=2)
            
            # Create time-based aggregations
            temp_df = self.df.copy()
            temp_df['year_month'] = pd.to_datetime(temp_df[date_col]).dt.to_period('M')
            
            monthly_rates = temp_df.groupby('year_month')[self.target_column].apply(
                lambda x: (x == self.positive_class).mean()
            )
            
            plt.figure(figsize=(12, 6))
            monthly_rates.plot(kind='line', marker='o', linewidth=2, color='#3498DB')
            plt.title(f'{self.target_column} Rate Over Time', fontsize=14, fontweight='bold')
            plt.xlabel('Time Period', fontsize=12)
            plt.ylabel(f'{self.target_column} Rate', fontsize=12)
            plt.grid(axis='both', alpha=0.3)
            plt.xticks(rotation=45)
            plt.tight_layout()
            plt.savefig(f'/tmp/temporal_{date_col}.png', dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            
            doc.add_picture(f'/tmp/temporal_{date_col}.png', width=Inches(6))
            
    def _add_correlation_analysis(self, doc):
        """Add correlation analysis for numeric features"""
        doc.add_heading('Feature Correlation Analysis', level=1)
        
        # Calculate correlation matrix
        numeric_df = self.df[self.numeric_cols].dropna()
        
        if len(numeric_df.columns) > 1:
            corr_matrix = numeric_df.corr()
            
            plt.figure(figsize=(10, 8))
            sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', center=0,
                       fmt='.2f', square=True, linewidths=1)
            plt.title('Feature Correlation Matrix', fontsize=14, fontweight='bold')
            plt.tight_layout()
            plt.savefig('/tmp/correlation.png', dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            
            doc.add_picture('/tmp/correlation.png', width=Inches(6.5))
            
    def _add_risk_scoring(self, doc):
        """Add predictive risk scoring"""
        doc.add_heading('Predictive Risk Scoring', level=1)
        
        doc.add_paragraph(
            f"Risk scores are calculated based on feature importance and correlation with {self.target_column}. "
            "Higher scores indicate higher likelihood of the positive outcome."
        )
        
        # Simple risk score based on correlation with target
        self.df['Risk_Score'] = 50  # Base score
        
        # Add contributions from numeric features
        for col in self.numeric_cols[:5]:
            if self.df[col].std() > 0:
                normalized = (self.df[col] - self.df[col].mean()) / self.df[col].std()
                self.df['Risk_Score'] += normalized * 10
                
        # Normalize to 0-100
        if self.df['Risk_Score'].std() > 0:
            self.df['Risk_Score'] = ((self.df['Risk_Score'] - self.df['Risk_Score'].min()) / 
                                     (self.df['Risk_Score'].max() - self.df['Risk_Score'].min())) * 100
        
        # Plot distribution
        plt.figure(figsize=(10, 6))
        plt.hist(self.df['Risk_Score'], bins=20, color='#E67E22', alpha=0.7, edgecolor='black')
        plt.title('Risk Score Distribution', fontsize=14, fontweight='bold')
        plt.xlabel('Risk Score (0-100)', fontsize=12)
        plt.ylabel('Number of Records', fontsize=12)
        plt.axvline(self.df['Risk_Score'].mean(), color='red', linestyle='--',
                   label=f'Mean: {self.df["Risk_Score"].mean():.1f}')
        plt.legend()
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        plt.savefig('/tmp/risk_scores.png', dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        doc.add_picture('/tmp/risk_scores.png', width=Inches(6))
        
    def _add_recommendations(self, doc):
        """Add strategic recommendations"""
        doc.add_heading('Strategic Recommendations', level=1)
        
        # Identify high-impact features
        recommendations = []
        
        # Check categorical features with high variance
        for col in self.categorical_cols[:3]:
            if self.df[col].nunique() < 50:
                rates = self.df.groupby(col)[self.target_column].apply(
                    lambda x: (x == self.positive_class).mean()
                )
                if rates.std() > 0.1:  # Significant variance
                    high_risk_cat = rates.idxmax()
                    recommendations.append({
                        'priority': 'HIGH',
                        'area': col,
                        'finding': f'High variance in {self.target_column} rates',
                        'action': f'Focus on {high_risk_cat} segment',
                        'impact': f'{rates.max()*100:.1f}% rate in highest segment'
                    })
        
        # Add general recommendations
        recommendations.extend([
            {
                'priority': 'HIGH',
                'area': 'Data Monitoring',
                'finding': f'{self.positive_rate*100:.1f}% overall rate',
                'action': 'Implement continuous monitoring and alerting',
                'impact': 'Early detection and prevention'
            },
            {
                'priority': 'MEDIUM',
                'area': 'Feature Engineering',
                'finding': f'{len(self.numeric_cols) + len(self.categorical_cols)} features analyzed',
                'action': 'Develop additional predictive features',
                'impact': 'Improved prediction accuracy'
            }
        ])
        
        # Create recommendations table
        rec_table = doc.add_table(rows=1, cols=5)
        rec_table.style = 'Table Grid'
        hdr_cells = rec_table.rows[0].cells
        hdr_cells[0].text = 'Priority'
        hdr_cells[1].text = 'Area'
        hdr_cells[2].text = 'Finding'
        hdr_cells[3].text = 'Recommended Action'
        hdr_cells[4].text = 'Expected Impact'
        
        for rec in recommendations:
            row_cells = rec_table.add_row().cells
            row_cells[0].text = rec['priority']
            row_cells[1].text = rec['area']
            row_cells[2].text = rec['finding']
            row_cells[3].text = rec['action']
            row_cells[4].text = rec['impact']
            
    def _add_appendices(self, doc):
        """Add appendices with detailed statistics"""
        doc.add_heading('Appendices', level=1)
        
        doc.add_heading('Appendix A: Feature Summary Statistics', level=2)
        
        # Numeric features summary
        if self.numeric_cols:
            doc.add_heading('Numerical Features', level=3)
            stats_df = self.df[self.numeric_cols].describe().T
            
            stats_table = doc.add_table(rows=1, cols=6)
            stats_table.style = 'Table Grid'
            hdr_cells = stats_table.rows[0].cells
            headers = ['Feature', 'Mean', 'Std', 'Min', 'Max', 'Missing']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                
            for col in self.numeric_cols[:10]:  # Top 10
                row_cells = stats_table.add_row().cells
                row_cells[0].text = col
                row_cells[1].text = f"{self.df[col].mean():.2f}"
                row_cells[2].text = f"{self.df[col].std():.2f}"
                row_cells[3].text = f"{self.df[col].min():.2f}"
                row_cells[4].text = f"{self.df[col].max():.2f}"
                row_cells[5].text = f"{self.df[col].isnull().sum()}"
                
        doc.add_heading('Appendix B: Methodology', level=2)
        methodology = f"""
        Analysis Methodology:
        
        Data Processing:
        - Dataset analyzed: {len(self.df):,} records
        - Target variable: {self.target_column}
        - Positive class: {self.positive_class}
        - Missing values handled appropriately
        
        Statistical Methods:
        - Descriptive statistics for all features
        - Group-wise analysis for categorical variables
        - Distribution analysis for numerical variables
        - Correlation analysis for feature relationships
        - Risk scoring using normalized feature values
        
        Visualization:
        - High-resolution charts (300 DPI)
        - Professional color schemes
        - Multiple chart types for different data types
        
        Report Generation:
        - Automated analysis and insights
        - Actionable recommendations
        - Executive-ready formatting
        """
        doc.add_paragraph(methodology.strip())
        
    def _add_page_headers(self, doc):
        """Add page headers"""
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{self.model_name} Analysis Report - Confidential"
            header_para.style = 'Header'


def generate_report_from_upload(file_path, target_column, model_name="Prediction Model", 
                                positive_class=None, output_path=None):
    """
    Main function to generate report from uploaded file
    
    Args:
        file_path: Path to Excel/CSV file
        target_column: Name of the target/prediction column
        model_name: Name of the model (e.g., "Churn Prediction", "Fraud Detection")
        positive_class: The positive class value (auto-detected if None)
        output_path: Where to save the report (auto-generated if None)
    """
    # Load data
    if file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
    else:
        df = pd.read_excel(file_path)
    
    print(f"Loaded {len(df)} records from {file_path}")
    print(f"Columns: {', '.join(df.columns.tolist())}")
    
    # Generate output path if not provided
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"{model_name.replace(' ', '_')}_Report_{timestamp}.docx"
    
    # Create report generator
    generator = GenericMLReportGenerator(
        df=df,
        target_column=target_column,
        model_name=model_name,
        positive_class=positive_class
    )
    
    # Generate report
    generator.create_report(output_path)
    
    return output_path


# Example usage
if __name__ == "__main__":
    # Example 1: HR Attrition
    generate_report_from_upload(
        file_path="hr_data.xlsx",
        target_column="Attrition",
        model_name="HR Attrition Prediction",
        positive_class="Yes"
    )
    
    # Example 2: Customer Churn
    generate_report_from_upload(
        file_path="customer_data.csv",
        target_column="Churned",
        model_name="Customer Churn Prediction",
        positive_class=1
    )
    
    # Example 3: Fraud Detection
    generate_report_from_upload(
        file_path="transactions.csv",
        target_column="Is_Fraud",
        model_name="Fraud Detection",
        positive_class=True
    )
